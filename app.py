import streamlit as st
from docx import Document
from datetime import datetime
import os
import platform
import subprocess
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
import fitz
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import tempfile
import uuid
import logging
import sys
from num2words import num2words
from google.cloud import firestore
import firebase_admin
from firebase_admin import credentials, firestore as admin_firestore, storage
import json
port = int(os.environ.get("PORT", 8080))  # Default to 8080
st.set_page_config(page_title="PDF Generator")
st.title("PDF Generator App")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)

logger = logging.getLogger("pdf_generator")
# Add this near the top of your main function or relevant document generator functions
def initialize_session_state():
    # Initialize all session state variables used in the app
    for key in ["nda_docx", "nda_pdf", "nda_docx_name", "nda_pdf_name", 
                "contract_docx", "contract_pdf", "contract_docx_name", "contract_pdf_name", 
                "hiring_docx", "hiring_pdf", "hiring_docx_name", "hiring_pdf_name",
                "invoice_docx", "invoice_pdf", "invoice_docx_name", "invoice_pdf_name"]:
        if key not in st.session_state:
            st.session_state[key] = None if "name" not in key else ""

# Call this function at the start of your main function
def convert_to_pdf(doc_path, pdf_path):
    doc_path = os.path.abspath(doc_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found at {doc_path}")

    # Use a temporary directory for intermediate files
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_pdf_path = os.path.join(temp_dir, "temp_output.pdf")

        # Step 1: Convert Word to PDF
        if platform.system() == "Windows":
            try:
                import comtypes.client
                import pythoncom
                pythoncom.CoInitialize()
                word = comtypes.client.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(doc_path)
                doc.SaveAs(temp_pdf_path, FileFormat=17)  # FileFormat=17 is for PDF
                doc.Close()
                word.Quit()
                pythoncom.CoUninitialize()
            except Exception as e:
                raise Exception(f"Error using COM on Windows: {e}")
        else:
            try:
                subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, doc_path],
                    check=True
                )
                # LibreOffice will save the PDF in the same folder as the DOCX with .pdf extension
                temp_pdf_path = os.path.join(temp_dir, os.path.basename(doc_path).replace('.docx', '.pdf'))
            except subprocess.CalledProcessError as e:
                raise Exception(f"Error using LibreOffice: {e}")

        # Step 2: Flatten the PDF (convert to image-based PDF)
        # Save it to the final location (pdf_path) instead of keeping it inside the temp folder
        # flatten_pdf(temp_pdf_path, pdf_path)
        import shutil
        shutil.copy(temp_pdf_path, pdf_path)

        # Optionally, confirm it was created
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Flattened PDF file was not saved correctly: {pdf_path}")

# Common Functions (unchanged)
def apply_formatting(run, font_name, font_size, bold=False):
    """Apply specific formatting to a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold



def apply_image_placeholder(doc, placeholder_key, image_file):
    """Replace a placeholder with an image in the Word document."""
    try:
        placeholder_found = False

        # Check inside tables first
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder_key in para.text:
                            para.clear()  # Clears text while preserving formatting
                            run = para.add_run()
                            run.add_picture(image_file, width=Inches(1.5), height=Inches(0.75))
                            placeholder_found = True

        # Check paragraphs outside tables
        for para in doc.paragraphs:
            if placeholder_key in para.text:
                para.clear()
                run = para.add_run()
                run.add_picture(image_file, width=Inches(1.2), height=Inches(0.75))
                placeholder_found = True

        if not placeholder_found:
            logger.warning(f"Placeholder '{placeholder_key}' not found in the document.")
        
        return doc

    except Exception as e:
        logger.error(f"Error inserting image: {e}", exc_info=True)
        return None  # Returning None to indicate failure

# Contract Generator
def replace_placeholders(doc, placeholders):
    """Replace placeholders in paragraphs and tables."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        # Apply bold formatting for specific placeholders
                        if key == "<<EndDate>>":
                            run.bold = True  # Apply bold formatting

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    # Apply bold formatting for specific placeholders
                                    if key == "<<EndDate>>":
                                        run.bold = True  # Apply bold formatting
    return doc


def replace_text_in_table(table, placeholders):
    """Iterate over all cells and paragraphs in a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders(paragraph, placeholders)


def edit_contract_template(template_path, output_path, placeholders):
    """Load a Word template, replace placeholders, and save to output."""
    doc = Document(template_path)
    # Replace in document body
    for para in doc.paragraphs:
        replace_placeholders(para, placeholders)
    # Replace in all tables
    for table in doc.tables:
        replace_text_in_table(table, placeholders)
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_contract():
    """Streamlit app to gather inputs and generate the contract document."""
    st.title("Contract Generator")

    # User inputs
    client_name = st.text_input("Enter Client Name:")
    company_name = st.text_input("Enter Company Name:")
    date_input = st.date_input("Enter Effective Date:", datetime.today())
    end_date_input = st.date_input("Enter Contract End Date:", datetime.today())
    address = st.text_area("Enter Address:")

    placeholders = {
        "<<ClientName>>": client_name,
        "<<CompanyName>>": company_name,
        "<<Date>>": date_input.strftime("%d-%m-%Y"),
        "<<StartDate>>": date_input.strftime("%d %B %Y"),
        "<<EndDate>>": end_date_input.strftime("%d %B %Y"),
        "<<Address>>": address,
    }

    template_name = "Contract Template.docx"
    
    if st.button("Generate Contract"):
        try:
            # Clear previous session state data
            for key in ["contract_docx", "contract_pdf", "contract_docx_name", "contract_pdf_name"]:
                if key in st.session_state:
                    st.session_state[key] = None if "name" not in key else ""

            # Define the hiring template file path
            template_path = os.path.join(os.getcwd(), template_name)
            
            # Verify the template exists
            if not os.path.exists(template_path):
                st.error(f"Template file not found: {template_path}")
                return

            
            safe_name = ''.join(c if c.isalnum() else '_' for c in client_name)
            
            # Save the contract to a temporary directory
            temp_dir = tempfile.gettempdir()
            docx_output_path = os.path.join(temp_dir, f"Contract_{safe_name}.docx")
            pdf_output_path = os.path.join(temp_dir, f"Contract_{safe_name}.pdf")

            # Edit the template and save the contract
            edit_contract_template(template_path, docx_output_path, placeholders)
            # st.info("DOCX file created successfully. Converting to PDF...")

            # Load the generated DOCX file into session state for download
            with open(docx_output_path, "rb") as docx_file:
                st.session_state.contract_docx = docx_file.read()
                st.session_state.contract_docx_name = f"Contract_{safe_name}.docx"

            # Convert DOCX to PDF with better error handling
            try:
                convert_to_pdf(docx_output_path, pdf_output_path)
                # st.info(f"PDF conversion completed. Checking result...")
                
                if os.path.exists(pdf_output_path):
                    with open(pdf_output_path, "rb") as pdf_file:
                        st.session_state.contract_pdf = pdf_file.read()
                        st.session_state.contract_pdf_name = f"Contract_{safe_name}.pdf"
                    # st.success("PDF created successfully!")
                else:
                    st.warning("PDF file not found after conversion attempt.")
            except Exception as pdf_err:
                st.error(f"PDF Conversion Error: {pdf_err}")
                # Still allow DOCX download even if PDF fails
                st.warning("PDF conversion failed, but DOCX is available for download.")

            # Display download buttons based on what's available
            col1, col2 = st.columns(2)
            
            with col1:
                if st.session_state.contract_docx:
                    st.download_button(
                        label="📥 Download Contract (Word)",
                        data=st.session_state.contract_docx,
                        file_name=st.session_state.contract_docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.warning("DOCX file not available for download.")
                    
            with col2:
                if st.session_state.contract_pdf:
                    st.download_button(
                        label="📥 Download Contract (PDF)",
                        data=st.session_state.contract_pdf,
                        file_name=st.session_state.contract_pdf_name,
                        mime="application/pdf"
                    )
                else:
                    st.warning("PDF file not available for download.")
                    
        except Exception as e:
            st.error(f"An error occurred: {e}")
            import traceback
            st.code(traceback.format_exc())

#NDA 
def replace_text_in_paragraph(paragraph, placeholders):
    """Replace placeholders in a paragraph, preserving formatting and optionally bolding specific runs."""
    # Combine all run texts
    full_text = ''.join(run.text for run in paragraph.runs)

    for key, value in placeholders.items():
        if key in full_text:
            before, sep, after = full_text.partition(key)
            # Remove existing runs
            for run in list(paragraph.runs)[::-1]:
                paragraph._p.remove(run._r)
            # Add text before placeholder
            if before:
                paragraph.add_run(before)
            # Add replaced text
            new_run = paragraph.add_run(value)
            # Uncomment the next lines if you want to bold any specific placeholders:
            # if key in ('<<ClientName>>', '<<Date>>'):
            #     new_run.font.bold = True
            # Add text after placeholder
            if after:
                paragraph.add_run(after)
            # Stop after first replacement in this paragraph
            return


def replace_text_in_table(table, placeholders):
    """Iterate through table cells and replace placeholders."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, placeholders)


def edit_nda_template(template_path, output_path, placeholders):
    """Load template, replace placeholders in body and tables, then save."""
    doc = Document(template_path)

    # Replace in document body
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, placeholders)

    # Replace in all tables
    for table in doc.tables:
        replace_text_in_table(table, placeholders)

    # Ensure target directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_nda():
    """Streamlit UI to collect inputs and generate the NDA document."""
    st.title("NDA Generator")

    client_name = st.text_input("Enter Client Name:")
    company_name = st.text_input("Enter Company Name:")
    date_input = st.date_input("Enter Date:", datetime.today())
    address = st.text_area("Enter Address:")

    placeholders = {
        "<<ClientName>>": client_name,
        "<<CompanyName>>": company_name,
        "<<Date>>": date_input.strftime("%d-%m-%Y"),
        "<<Address>>": address,
    }

    template_name = "NDA Template.docx"
    temp_dir = tempfile.gettempdir()


    if st.button("Generate NDA"):
        try:
            # Clear previous session state data
            for key in ["nda_docx", "nda_pdf", "nda_docx_name", "nda_pdf_name"]:
                if key in st.session_state:
                    st.session_state[key] = None if "name" not in key else ""

            # Define the hiring template file path
            template_path = os.path.join(os.getcwd(), template_name)
            
            # Verify the template exists
            if not os.path.exists(template_path):
                st.error(f"Template file not found: {template_path}")
                return

            
            safe_name = ''.join(c if c.isalnum() else '_' for c in client_name)
            
            # Save the hiring contract to a temporary directory
            temp_dir = tempfile.gettempdir()
            docx_output_path = os.path.join(temp_dir, f"NDA_{safe_name}.docx")
            pdf_output_path = os.path.join(temp_dir, f"NDA_{safe_name}.pdf")

            # Edit the hiring template and save the contract
            edit_nda_template(template_path, docx_output_path, placeholders)
            # st.info("DOCX file created successfully. Converting to PDF...")

            # Load the generated DOCX file into session state for download
            with open(docx_output_path, "rb") as docx_file:
                st.session_state.nda_docx = docx_file.read()
                st.session_state.nda_docx_name = f"NDA_{safe_name}.docx"

            # Convert DOCX to PDF with better error handling
            try:
                convert_to_pdf(docx_output_path, pdf_output_path)
                # st.info(f"PDF conversion completed. Checking result...")
                
                if os.path.exists(pdf_output_path):
                    with open(pdf_output_path, "rb") as pdf_file:
                        st.session_state.nda_pdf = pdf_file.read()
                        st.session_state.nda_pdf_name =f"NDA_{safe_name}.pdf"
                    # st.success("PDF created successfully!")
                else:
                    st.warning("PDF file not found after conversion attempt.")
            except Exception as pdf_err:
                st.error(f"PDF Conversion Error: {pdf_err}")
                # Still allow DOCX download even if PDF fails
                st.warning("PDF conversion failed, but DOCX is available for download.")

            # Display download buttons based on what's available
            col1, col2 = st.columns(2)
            
            with col1:
                if st.session_state.nda_docx:
                    st.download_button(
                        label="📥 Download NDA(Word)",
                        data=st.session_state.nda_docx,
                        file_name=st.session_state.nda_docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.warning("DOCX file not available for download.")
                    
            with col2:
                if st.session_state.nda_pdf:
                    st.download_button(
                        label="📥 Download NDA(PDF)",
                        data=st.session_state.nda_pdf,
                        file_name=st.session_state.nda_pdf_name,
                        mime="application/pdf"
                    )
                else:
                    st.warning("PDF file not available for download.")
                    
        except Exception as e:
            st.error(f"An error occurred: {e}")
            import traceback
            st.code(traceback.format_exc())

# Hiring Contract
def replace_text_in_paragraph(paragraph, placeholders):
    """Replace placeholders in paragraph text."""
    if not paragraph.runs:
        return
        
    full_text = "".join(run.text for run in paragraph.runs)
    
    for key, value in placeholders.items():
        if key in full_text:
            full_text = full_text.replace(key, value)
    
    # Clear all runs
    for i in range(len(paragraph.runs) - 1, 0, -1):
        p = paragraph._p
        p.remove(paragraph.runs[i]._r)
    
    # Assign replaced full text to the first run
    if paragraph.runs:
        paragraph.runs[0].text = full_text

# Function to edit the Hiring template and replace placeholders
def edit_hiring_template(template_path, output_path, placeholders):
    """Edit hiring contract template replacing placeholders."""
    doc = Document(template_path)

    def replace_text_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, placeholders)

    for para in doc.paragraphs:
        replace_text_in_paragraph(para, placeholders)

    for table in doc.tables:
        replace_text_in_table(table)

    doc.save(output_path)
    return output_path

def generate_hiring_contract():
    """Streamlit UI for generating hiring contracts."""
    # Initialize session state for DOCX and PDF
    for key in ["hiring_docx", "hiring_pdf", "hiring_docx_name", "hiring_pdf_name"]:
        if key not in st.session_state:
            st.session_state[key] = None if "name" not in key else ""

    st.title("Hiring Contract Generator")

    # Collect inputs
    Employee_name = st.text_input("Enter Employee Name:")
    Role = st.text_input("Enter Role:")
    date = st.date_input("Enter Today Date:", datetime.today())
    Starting_Date = st.date_input("Enter the starting date: ")
    Stipend = st.number_input("Enter the Stipend:")
    Working_hours = st.number_input("Enter the total working hours:")
    Internship_duration = st.number_input("Enter the internship duration (in months):")
    First_Pay_Cheque_Date = st.date_input("Enter the First Pay Cheque Date:")

    placeholders = {
        "<<Date>>": date.strftime("%d %B, %Y"),
        "<<Name>>": Employee_name,
        "<<Role>>": Role,
        "<<Starting Date>>": Starting_Date.strftime("%d %B, %Y"),
        "<<Internship Duration>>": str(int(Internship_duration)),
        "<<First Pay>>": First_Pay_Cheque_Date.strftime("%d %B, %Y"),
        "<<Stipend>>": str(Stipend),
        "<<Working Hours>>": str(int(Working_hours))
    }

    template_name = "Hiring Contract.docx"
    
    if st.button("Generate Hiring Contract"):
        try:
            # Clear previous session state data
            for key in ["hiring_docx", "hiring_pdf", "hiring_docx_name", "hiring_pdf_name"]:
                if key in st.session_state:
                    st.session_state[key] = None if "name" not in key else ""

            # Define the hiring template file path
            template_path = os.path.join(os.getcwd(), template_name)
            
            # Verify the template exists
            if not os.path.exists(template_path):
                st.error(f"Template file not found: {template_path}")
                return

            
            safe_name = ''.join(c if c.isalnum() else '_' for c in Employee_name)
            
            # Save the hiring contract to a temporary directory
            temp_dir = tempfile.gettempdir()
            docx_output_path = os.path.join(temp_dir, f"Hiring_{safe_name}.docx")
            pdf_output_path = os.path.join(temp_dir, f"Hiring_{safe_name}.pdf")

            # Edit the hiring template and save the contract
            edit_hiring_template(template_path, docx_output_path, placeholders)
            # st.info("DOCX file created successfully. Converting to PDF...")

            # Load the generated DOCX file into session state for download
            with open(docx_output_path, "rb") as docx_file:
                st.session_state.hiring_docx = docx_file.read()
                st.session_state.hiring_docx_name = f"Hiring_{safe_name}.docx"

            # Convert DOCX to PDF with better error handling
            try:
                convert_to_pdf(docx_output_path, pdf_output_path)
                # st.info(f"PDF conversion completed. Checking result...")
                
                if os.path.exists(pdf_output_path):
                    with open(pdf_output_path, "rb") as pdf_file:
                        st.session_state.hiring_pdf = pdf_file.read()
                        st.session_state.hiring_pdf_name = f"Hiring_{safe_name}.pdf"
                    # st.success("PDF created successfully!")
                else:
                    st.warning("PDF file not found after conversion attempt.")
            except Exception as pdf_err:
                st.error(f"PDF Conversion Error: {pdf_err}")
                # Still allow DOCX download even if PDF fails
                st.warning("PDF conversion failed, but DOCX is available for download.")

            # Display download buttons based on what's available
            col1, col2 = st.columns(2)
            
            with col1:
                if st.session_state.hiring_docx:
                    st.download_button(
                        label="📥 Download Hiring Contract (Word)",
                        data=st.session_state.hiring_docx,
                        file_name=st.session_state.hiring_docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.warning("DOCX file not available for download.")
                    
            with col2:
                if st.session_state.hiring_pdf:
                    st.download_button(
                        label="📥 Download Hiring Contract (PDF)",
                        data=st.session_state.hiring_pdf,
                        file_name=st.session_state.hiring_pdf_name,
                        mime="application/pdf"
                    )
                else:
                    st.warning("PDF file not available for download.")
                    
        except Exception as e:
            st.error(f"An error occurred: {e}")
            import traceback
            st.code(traceback.format_exc())


def format_price(amount, currency):
    """Format price based on currency."""
    formatted_price = f"{amount:,.2f}"
    return f"{currency} {formatted_price}" if currency == "USD" else f"Rs. {formatted_price}"

# Function to format percentages
def format_percentage(value):
    """Format percentage without decimals."""
    return f"{int(value)}%"

# Function to get the next invoice number
def get_next_invoice_number():
    """Fetch and increment the invoice number."""
    invoice_file = "invoice_counter.txt"
    if not os.path.exists(invoice_file):
        with open(invoice_file, "w") as file:
            file.write("1000")  # Starting invoice number
    try:
        with open(invoice_file, "r") as file:
            current_invoice = int(file.read().strip())
    except ValueError:
        current_invoice = 1000
    next_invoice = current_invoice + 1
    with open(invoice_file, "w") as file:
        file.write(str(next_invoice))
    return current_invoice

# Function to convert amount to words
def amount_to_words(amount):
    """Convert amount to words without currency formatting."""
    try:
        words = num2words(amount, lang='en').replace(',', '').title()
        return words
    except Exception as e:
        logger.error(f"Error converting amount to words: {e}", exc_info=True)
        return f"[Error: Unable to convert {amount} to words]"

# Function to replace placeholders in DOCX
def replace_placeholders(doc, placeholders):
    """Replace placeholders in paragraphs and tables."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        # Apply bold formatting for specific placeholders
                        if key.startswith("<<Price") or key.startswith("<<Total") or key == "<<Amt to Word>>":
                            run.bold = True  # Apply bold formatting

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    # Apply bold formatting for specific placeholders
                                    if key.startswith("<<Price") or key.startswith("<<Total") or key == "<<Amt to Word>>":
                                        run.bold = True  # Apply bold formatting
    return doc

# Function to edit invoice template and save the result
def edit_invoice_template(template_name, output_path, placeholders):
    """Edit an invoice template and save the result."""
    try:
        doc = Document(template_name)
        replace_placeholders(doc, placeholders)
        doc.save(output_path)
        return output_path
    except Exception as e:
        logger.error(f"Error editing invoice template: {e}", exc_info=True)
        raise Exception(f"Error editing invoice template: {e}")


def generate_invoice():
    """Streamlit app for generating invoices."""
    st.title("Invoice Generator")
    # Select region
    region = st.selectbox("Select Region", ["INR", "USD"])

    # Set payment options dynamically
    payment_options = ["1 Payment", "3 EMI", "5 EMI"] if region == "INR" else ["3 EMI", "5 EMI"]

    # Input Fields
    client_name = st.text_input("Client Name")
    client_address = st.text_input("Client Address")
    project_name = st.text_input("Project Name")
    phone_number = st.text_input("Phone Number")
    gst_number = st.text_input("GST Number")
    base_amount = st.number_input("Base Amount (excluding GST)", min_value=0.0, format="%.2f")
    payment_option = st.selectbox("Payment Option", payment_options)
    invoice_date = st.date_input("Invoice Date", value=datetime.today())
    formatted_date = invoice_date.strftime("%d-%m-%Y")

    # Calculate GST and total amount
    gst_amount = round(base_amount * 0.18)
    total_amount = base_amount + gst_amount

    # Prepare placeholders for template
    placeholders = {
        "<<Client Name>>": client_name,
        "<<Client Address>>": client_address,
        "<<GST Number>>": gst_number,
        "<<Client Email>>": client_address,
        "<<Project Name>>": project_name,
        "<<Mobile Number>>": phone_number,
        "<<Date>>": formatted_date,
        "<<Amt to word>>": amount_to_words(int(total_amount)),
    }
# Select the correct template based on payment option
    if payment_option == "1 Payment":
        template_name = f"Invoice Template - {region} - 1 Payment 1.docx"
        placeholders.update({
            "<<Price 1>>": format_price(base_amount, region),
            "<<Price 2>>": format_price(gst_amount, region),
            "<<Price 3>>": format_price(total_amount, region),
            "<<Total 1>>": format_price(total_amount, region),
        })

    elif payment_option == "3 EMI":
        template_name = f"Invoice Template - {region} - 3 EMI Payment Schedule 1.docx"
        p1 = round(total_amount * 0.30)
        p2 = round(total_amount * 0.40)
        p3 = total_amount - (p1 + p2)
        placeholders.update({
            "<<Price 1>>": format_price(p1, region),
            "<<Price 2>>": format_price(p2, region),
            "<<Price 3>>": format_price(p3, region),
            "<<Price 4>>": format_price(gst_amount, region),
            "<<Price 5>>": format_price(total_amount, region),
            "<<Price 6>>": format_price(p1, region),
            "<<Price 7>>": format_price(p2, region),
            "<<Price 8>>": format_price(p3, region),
        })

    elif payment_option == "5 EMI":
        template_name = f"Invoice Template - {region} - 5 EMI Payment Schedule 1.docx"
        p1 = round(total_amount * 0.20)
        p2 = round(total_amount * 0.20)
        p3 = round(total_amount * 0.20)
        p4 = round(total_amount * 0.20)
        p5 = total_amount - (p1 + p2 + p3 + p4)
        placeholders.update({
            "<<Price 1>>": format_price(p1, region),
            "<<Price 2>>": format_price(p2, region),
            "<<Price 3>>": format_price(p3, region),
            "<<Price 4>>": format_price(p4, region),
            "<<Price 5>>": format_price(p5, region),
            "<<Price 6>>": format_price(p1, region),
            "<<Price 7>>": format_price(p2, region),
            "<<Price 8>>": format_price(p3, region),
            "<<Price 9>>": format_price(p4, region),
            "<<Price 10>>": format_price(p5, region),
            "<<Total 1>>": format_price(total_amount, region),
        })

    # Generate Invoice Button
    if st.button("Generate Invoice"):
        try:
            for key in ["invoice_docx","invoice_pdf","invoice_docx_name","invoice_pdf_name"]:
                if key in st.session_state:
                    st.session_state[key] =None if "name" not in key else ""

            # Generate the next invoice number
            invoice_number = get_next_invoice_number()
            placeholders["<<Invoice>>"] = str(invoice_number)
            placeholders["<<Invoice No>>"] = str(invoice_number)

            # Define the invoice template file path
            template_path = os.path.join(os.getcwd(), template_name)

            if not os.path.exists(template_path):
                st.error(f"Template file not found: {template_path}")
                return

            # Save the invoice to a temporary directory
            temp_dir = tempfile.gettempdir()
            sanitized_client_name = "".join([c if c.isalnum() or c.isspace() else "_" for c in client_name])
            docx_output_path = os.path.join(temp_dir, f"Invoice_{sanitized_client_name}_{invoice_number}.docx")
            pdf_output_path = os.path.join(temp_dir, f"Invoice_{sanitized_client_name}_{invoice_number}.pdf")

            # Edit the template and save the invoice
            edit_invoice_template(template_path, docx_output_path, placeholders)
            
           #Save the file to session
            with open(docx_output_path, "rb") as docx_file:
                st.session_state.invoice_docx = docx_file.read()
                st.session_state.invoice_docx_name = f"Invoice_{sanitized_client_name}_{invoice_number}.docx"
            # Convert DOCX to PDF with better error handling
            try:
                convert_to_pdf(docx_output_path, pdf_output_path)
                st.info(f"PDF conversion completed. Checking result...")
                
                if os.path.exists(pdf_output_path):
                    with open(pdf_output_path, "rb") as pdf_file:
                        st.session_state.invoice_pdf = pdf_file.read()
                        st.session_state.invoice_pdf_name = f"Invoice_{sanitized_client_name}_{invoice_number}.pdf"
                    st.success("PDF created successfully!")
                else:
                    st.warning("PDF file not found after conversion attempt.")
            except Exception as pdf_err:
                st.error(f"PDF Conversion Error: {pdf_err}")
                # Still allow DOCX download even if PDF fails
                st.warning("PDF conversion failed, but DOCX is available for download.")

            # Display download buttons based on what's available
            col1, col2 = st.columns(2)
            
            with col1:
                if st.session_state.invoice_docx:
                    st.download_button(
                        label="📥 Download Invoice (Word)",
                        data=st.session_state.invoice_docx,
                        file_name=st.session_state.invoice_docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.warning("DOCX file not available for download.")
                    
            with col2:
                if st.session_state.invoice_pdf:
                    st.download_button(
                        label="📥 Download Invoice (PDF)",
                        data=st.session_state.invoice_pdf,
                        file_name=st.session_state.invoice_pdf_name,
                        mime="application/pdf"
                    )
                else:
                    st.warning("PDF file not available for download.")
                    
        except Exception as e:
            st.error(f"An error occurred: {e}")
            import traceback
            st.code(traceback.format_exc())


#Firebase operations

firebase_info = dict(st.secrets["FIREBASE"])

# Get the bucket name
bucket_name = "hv-technologies.appspot.com"  
# Initialize Firebase only if not already initialized
if not firebase_admin._apps:
    cred = credentials.Certificate(firebase_info)
    firebase_admin.initialize_app(cred, {
        'storageBucket': bucket_name
    })

# Now you can safely get the bucket
bucket = storage.bucket()

# ========== UTILITY FUNCTIONS ==========

def get_db():
    """Returns Firestore client."""
    return firestore.Client.from_service_account_json("hv-technologies-firebase-adminsdk.json")


# ========== CREATE / UPLOAD ==========

def upload_to_firebase(uploaded_file, name):
    """Uploads PDF to Firebase Storage and saves metadata to Firestore."""
    blob = bucket.blob(f"uploaded_docs/{uploaded_file.name}")
    blob.upload_from_file(uploaded_file, content_type=uploaded_file.type)
    blob.make_public()

    link = blob.public_url
    db.collection("ProposalPDFPage2").document().set({
        "name": name,
        "link": link
    })

    st.success("File uploaded successfully!")
    st.markdown(f"[Click to View]({link})")


# ========== READ / SHOW ==========

def show_documents():
    st.markdown("### Uploaded Documents")
    docs = list(db.collection("ProposalPDFPage2").stream())

    for idx, doc in enumerate(docs, 1):
        doc_data = doc.to_dict()
        name = doc_data.get("name", "No Name")
        link = doc_data.get("link", "").strip()

        st.markdown(f"**{idx}. {name}**")
        if link.endswith(".pdf"):
            st.markdown(f"""
                <iframe src="{link}" width="100%" height="500px" style="border:1px solid #ccc;"></iframe>
            """, unsafe_allow_html=True)
        elif link:
            st.markdown(f"🔗 [View Document]({link})")
        else:
            st.warning("_No link available_")


# ========== UPDATE / DELETE ==========

def update_document(doc_id, new_name, new_link):
    db.collection("ProposalPDFPage2").document(doc_id).update({
        "name": new_name,
        "link": new_link.strip()
    })
    st.success("Document updated successfully!")

def delete_document(doc_id):
    db.collection("ProposalPDFPage2").document(doc_id).delete()
    st.success("Document deleted successfully!")

def manage_documents():
    docs = list(db.collection("ProposalPDFPage2").stream())

    if not docs:
        st.info("No documents found.")
        return

    doc_options = {f"{doc.to_dict().get('name', 'Unnamed')} ({doc.id})": doc.id for doc in docs}
    selected_label = st.selectbox("Select a document to update or delete", list(doc_options.keys()))
    selected_id = doc_options[selected_label]
    selected_data = db.collection("ProposalPDFPage2").document(selected_id).get().to_dict()

    st.markdown("### Update Document")
    with st.form("update_form"):
        updated_name = st.text_input("Updated Name", selected_data.get("name", ""))
        updated_link = st.text_input("Updated Link", selected_data.get("link", ""))
        update_btn = st.form_submit_button("Update")
        if update_btn:
            update_document(selected_id, updated_name, updated_link)

    st.markdown("###  Delete Document")
    if st.button("Delete This Document"):
        delete_document(selected_id)


# ========== STREAMLIT MAIN FUNCTION ==========


def main():
    #t.set_page_config(page_title="Document Generator & Firebase Manager", layout="wide")
    st.sidebar.title("📂 Application Menu")
    
    section = st.sidebar.radio("Choose Section", ["📄 Document Generators", "🔥 Firebase CRUD Operations"])

    if section == "📄 Document Generators":
        doc_choice = st.sidebar.radio("Select Document Type", ["NDA", "Contract", "Hiring Contract", "Invoice"])
        
        if doc_choice == "NDA":
            generate_nda()
        elif doc_choice == "Contract":
            generate_contract()
        elif doc_choice == "Hiring Contract":
            generate_hiring_contract()
        elif doc_choice == "Invoice":
            generate_invoice()

    elif section == "🔥 Firebase CRUD Operations":
        crud_choice = st.sidebar.radio("Choose Operation", ["Upload Document", "View Documents", "Update/Delete Document"])
        
        if crud_choice == "Upload Document":
            st.subheader("📤 Upload Document")
            with st.form("upload_form"):
                name = st.text_input("Enter Document Name")
                uploaded_file = st.file_uploader("Choose a file", type=["pdf"])
                submit_btn = st.form_submit_button("Upload")

                if submit_btn and uploaded_file and name:
                    upload_to_firebase(uploaded_file, name)

        elif crud_choice == "View Documents":
            show_documents()

        elif crud_choice == "Update/Delete Document":
            manage_documents()

if __name__ == "__main__":
    main()