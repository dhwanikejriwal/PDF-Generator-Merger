import streamlit as st
from docx import Document
from datetime import datetime
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import tempfile
from num2words import num2words
import uuid
from docx2pdf import convert  # For DOCX to PDF conversion
import pythoncom  # For COM initialization

# Common Functions (unchanged)
def apply_formatting(run, font_name, font_size, bold=False):
    """Apply specific formatting to a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold

def replace_and_format(doc, placeholders, font_name, font_size, option):
    """Replace placeholders and apply formatting."""
    for para in doc.paragraphs:
        if para.text:
            for key, value in placeholders.items():
                if key in para.text:
                    runs = para.runs
                    for run in runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
                            if para == doc.paragraphs[0]:
                                apply_formatting(run, font_name, font_size, bold=True)
                        else:
                            run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    for key, value in placeholders.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if key == "<<Address>>" else WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    apply_formatting(run, "Times New Roman", 11 if option == "NDA" else 12)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def edit_word_template(template_path, output_path, placeholders, font_name, font_size, option):
    """Edit Word document and apply formatting."""
    try:
        doc = Document(template_path)
        replace_and_format(doc, placeholders, font_name, font_size, option)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing Word template: {e}")

def apply_image_placeholder(doc, placeholder_key, image_file):
    """Replace a placeholder with an image in the Word document."""
    try:
        placeholder_found = False

        # Check inside tables first
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder_key in para.text:
                            para.clear()  # Clears text while preserving formatting
                            run = para.add_run()
                            run.add_picture(image_file, width=Inches(1.5), height=Inches(0.75))
                            placeholder_found = True

        # Check paragraphs outside tables
        for para in doc.paragraphs:
            if placeholder_key in para.text:
                para.clear()
                run = para.add_run()
                run.add_picture(image_file, width=Inches(1.2), height=Inches(0.75))
                placeholder_found = True

        if not placeholder_found:
            raise ValueError(f"Placeholder '{placeholder_key}' not found in the document.")
        
        return doc

    except Exception as e:
        print(f"Error inserting image: {e}")
        return None  # Returning None to indicate failure

# Contract/NDA Generator
def generate_document(option):
    """Streamlit UI for generating NDA or Contract documents."""
    st.title("Document Generator")

    base_dir = os.path.abspath(os.path.dirname(__file__))
    template_paths = {
        "NDA": os.path.join(base_dir, "NDA Template.docx"),
        "Contract": os.path.join(base_dir, "Contract Template.docx"),
    }

    client_name = st.text_input("Enter Client Name:")
    company_name = st.text_input("Enter Company Name:")
    address = st.text_area("Enter Address:")
    date_field = st.date_input("Enter Date:", datetime.today())
    # signature_file = st.file_uploader("Upload E-Signature (PNG or JPEG)", type=["png", "jpg", "jpeg"])

    placeholders = {
        "ClientName": client_name,
        "CompanyName": company_name,
        "Address": address,
        "Date": date_field.strftime("%d-%m-%Y"),
        "Date,": date_field.strftime("%d-%m-%Y"),
        "ContractEndDate": date_field.replace(year=1).strftime("%d-%m-%Y"),
        # "<< Date (Signature) >>": date_field.strftime("%d-%m-%Y"),
    }

    if st.button(f"Generate {option}"):
        try:
            # Clear previous session state data
            if 'doc_data' in st.session_state:
                del st.session_state.doc_data
                del st.session_state.pdf_data
                del st.session_state.filenames

            formatted_date = date_field.strftime("%d %b %Y")
            unique_id = str(uuid.uuid4())[:8]
            doc_filename = f"{option} - {client_name} {formatted_date} - {unique_id}.docx"
            pdf_filename = doc_filename.replace(".docx", ".pdf")

            # Create temporary files
            with tempfile.TemporaryDirectory() as temp_dir:
                doc_path = os.path.join(temp_dir, doc_filename)
                pdf_path = os.path.join(temp_dir, pdf_filename)

                # Generate DOCX
                font_size = 11 if option == "NDA" else 12
                doc = Document(template_paths[option])
                replace_and_format(doc, placeholders, "Times New Roman", font_size, option)

                # if signature_file:
                #     signature_path = os.path.join(temp_dir, "signature.png")
                #     with open(signature_path, "wb") as f:
                #         f.write(signature_file.getbuffer())
                #     doc = apply_image_placeholder(doc, "<<Signature>>", signature_path)

                doc.save(doc_path)

                # Convert to PDF
                pythoncom.CoInitialize()
                try:
                    convert(doc_path, pdf_path)
                finally:
                    pythoncom.CoUninitialize()

                # Store file data in session state
                with open(doc_path, "rb") as doc_file:
                    st.session_state.doc_data = doc_file.read()
                with open(pdf_path, "rb") as pdf_file:
                    st.session_state.pdf_data = pdf_file.read()
                
                st.session_state.filenames = {
                    "doc": doc_filename,
                    "pdf": pdf_filename
                }

            st.success(f"{option} Document Generated Successfully!")

        except Exception as e:
            st.error(f"An error occurred: {e}")

    # Display download buttons if data exists
    if 'doc_data' in st.session_state and 'pdf_data' in st.session_state:
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="Download Document (Word)",
                data=st.session_state.doc_data,
                file_name=st.session_state.filenames["doc"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_word"
            )
        with col2:
            st.download_button(
                label="Download Document (PDF)",
                data=st.session_state.pdf_data,
                file_name=st.session_state.filenames["pdf"],
                mime="application/pdf",
                key="download_pdf"
            )

# Hiring Contract Generator


def replace_placeholders(doc, placeholders):
    """Replace placeholders in paragraphs and tables."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    
                                    
    return doc

def edit_hiring_template(template_name, output_path, placeholders):
    """Edit an hiring contract template and save the result."""
    try:
        doc = Document(template_name)
        replace_placeholders(doc, placeholders)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing hiring contract template: {e}")
def generate_hiring_contract():
    if "hiring_docx" not in st.session_state:
        st.session_state.hiring_docx = None
    if "hiring_pdf" not in st.session_state:
        st.session_state.hiring_pdf = None
    if "hiring_docx_name" not in st.session_state:
        st.session_state.hiring_docx_name = ""
    if "hiring_pdf_name" not in st.session_state:
        st.session_state.hiring_pdf_name = ""
    
    st.title("Hiring Contract Generator")
    base_dir = os.path.abspath(os.path.dirname(__file__))
    template_paths = os.path.join(base_dir, "Hiring Contract.docx")

    Employee_name = st.text_input("Enter Employee Name:")
    Role = st.text_input("Enter Role:")
    date = st.date_input("Enter Date:", datetime.today())
    Starting_Date = st.date_input("Enter the starting date: ")
    Stipend = st.number_input("Enter the Stipend:")
    Working_hours = st.number_input("Enter the total working hours:")
    Internship_duration = st.number_input("Enter the internship duration:")
    First_Pay_Cheque_Date = st.date_input("Enter the First Pay Cheque Date:")

    placeholders = {
        "<<Today’s Date>>": date.strftime("%d-%m-%Y"),
        "<<Name>>": Employee_name,
        "<<Role>>": Role,
        "<<Starting Date>>": Starting_Date.strftime("%d-%m-%Y"),
        "<<Stipend>>": str(Stipend),
        "<<Working Hours>>": str(Working_hours),
        "<<Internship Duration months>>": str(Internship_duration),
        "<<First Pay Cheque Date>>": First_Pay_Cheque_Date.strftime("%d-%m-%Y")
    }

    if st.button("Generate Hiring Contract"):
        try:
            # Clear previous session state data
            if 'hiring_docx' in st.session_state:
                del st.session_state.hiring_docx
                del st.session_state.hiring_pdf
                del st.session_state.hiring_docx_name
                del st.session_state.hiring_pdf_name

            # Define the hiring contract template file path
            template_path = os.path.join(os.getcwd(), template_paths)

            # Save the hiring to a temporary directory
            temp_dir = tempfile.gettempdir()
            docx_output_path = os.path.join(temp_dir, f"Hiring Contract_{Employee_name}.docx")

            # Edit the template and save the hiring
            edit_hiring_template(template_path, docx_output_path, placeholders)

            # Check if the DOCX file was created successfully
            if not os.path.exists(docx_output_path):
                raise Exception(f"Failed to generate DOCX file at {docx_output_path}")

            # Read DOCX file bytes into session state
            with open(docx_output_path, "rb") as file:
                st.session_state.hiring_docx = file.read()
            st.session_state.hiring_docx_name = f"Hiring Contract_{Employee_name}.docx"

            # Generate PDF and store in session state
            pdf_output_path = os.path.join(temp_dir, f"Hiring Contract_{Employee_name}.pdf")
            pythoncom.CoInitialize()  # Initialize COM library
            try:
                convert(docx_output_path, pdf_output_path)
                # Check if the PDF file was created successfully
                if not os.path.exists(pdf_output_path):
                    raise Exception(f"Failed to generate PDF file at {pdf_output_path}")

                with open(pdf_output_path, "rb") as file:
                    st.session_state.hiring_pdf = file.read()
                st.session_state.hiring_pdf_name = f"Hiring Contract_{Employee_name}.pdf"
            finally:
                pythoncom.CoUninitialize()

            st.success(f"✅ Hiring Contract generated successfully!")

        except Exception as e:
            st.error(f"An error occurred: {e}")

    # Display download buttons if data exists in session state
    if 'hiring_docx' in st.session_state and 'hiring_pdf' in st.session_state:
        col1, col2 = st.columns(2)
        with col1:
            if st.session_state.hiring_docx:
                st.download_button(
                    label="📥 Download Hiring Contract (Word)",
                    data=st.session_state.hiring_docx,
                    file_name=st.session_state.hiring_docx_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.warning("DOCX file not available.")

        with col2:
            if st.session_state.hiring_pdf:
                st.download_button(
                    label="📥 Download Hiring Contract (PDF)",
                    data=st.session_state.hiring_pdf,
                    file_name=st.session_state.hiring_pdf_name,
                    mime="application/pdf"
                )
            else:
                st.warning("PDF file not available.")



# Invoice Generator (unchanged)
def format_price(amount, currency):
    """Format price based on currency."""
    formatted_price = f"{amount:,.2f}"
    return f"{currency} {formatted_price}" if currency == "USD" else f"Rs. {formatted_price}"

def format_percentage(value):
    """Format percentage without decimals."""
    return f"{int(value)}%"

def get_next_invoice_number():
    """Fetch and increment the invoice number."""
    invoice_file = "invoice_counter.txt"
    if not os.path.exists(invoice_file):
        with open(invoice_file, "w") as file:
            file.write("1000")  # Starting invoice number
    try:
        with open(invoice_file, "r") as file:
            current_invoice = int(file.read().strip())
    except ValueError:
        current_invoice = 1000
    next_invoice = current_invoice + 1
    with open(invoice_file, "w") as file:
        file.write(str(next_invoice))
    return current_invoice

def amount_to_words(amount):
    """Convert amount to words without currency formatting."""
    words = num2words(amount, lang='en').replace(',', '').title()
    
    return words

def replace_placeholders(doc, placeholders):
    """Replace placeholders in paragraphs and tables."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        # Apply bold formatting for specific placeholders
                        if key.startswith("<<Price") or key.startswith("<<Total") or key == "<<Amt to Word>>":
                            run.bold = True  # Apply bold formatting

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    # Apply bold formatting for specific placeholders
                                    if key.startswith("<<Price") or key.startswith("<<Total") or key == "<<Amt to Word>>":
                                        run.bold = True  # Apply bold formatting
    return doc

def edit_invoice_template(template_name, output_path, placeholders):
    """Edit an invoice template and save the result."""
    try:
        doc = Document(template_name)
        replace_placeholders(doc, placeholders)
        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Error editing invoice template: {e}")
def generate_invoice():
    """Streamlit app for generating invoices."""
    st.title("Invoice Generator")
    # Select region
    region = st.selectbox("Select Region", ["INR", "USD"])


    # Set payment options dynamically
    payment_options = ["1 Payment", "3 EMI", "5 EMI"] if region == "INR" else ["3 EMI", "5 EMI"]


    # Input Fields
    client_name = st.text_input("Client Name")
    client_address = st.text_input("Client Address")
    project_name = st.text_input("Project Name")
    phone_number = st.text_input("Phone Number")
    gst_number = st.text_input("GST Number")
    base_amount = st.number_input("Base Amount (excluding GST)", min_value=0.0, format="%.2f")
    payment_option = st.selectbox("Payment Option", payment_options)
    invoice_date = st.date_input("Invoice Date", value=datetime.today())
    formatted_date = invoice_date.strftime("%d-%m-%Y")

    # Calculate GST and total amount
    gst_amount = round(base_amount * 0.18)
    total_amount = base_amount + gst_amount

    # Prepare placeholders for template
    placeholders = {
        "<<Client Name>>": client_name,
        "<<Client Address>>": client_address,
        "<<GST Number>>": gst_number,
        "<<Client Email>>": client_address,
        "<<Project Name>>": project_name,
        "<<Mobile Number>>": phone_number,
        "<<Date>>": formatted_date,
        "<<Amt to word>>": amount_to_words(int(total_amount)),
    }

    # Select the correct template based on payment option
    if payment_option == "1 Payment":
        template_name = f"Invoice Template - {region} - 1 Payment 1.docx"
        placeholders.update({
            "<<Price 1>>": format_price(base_amount, region),
            "<<Price 2>>": format_price(gst_amount, region),
            "<<Price 3>>": format_price(total_amount, region),
            "<<Total 1>>": format_price(total_amount, region),
        })

    elif payment_option == "3 EMI":
        template_name = f"Invoice Template - {region} - 3 EMI Payment Schedule 1.docx"
        p1 = round(total_amount * 0.30)
        p2 = round(total_amount * 0.40)
        p3 = total_amount - (p1 + p2)
        placeholders.update({
            "<<Price 1>>": format_price(p1, region),
            "<<Price 2>>": format_price(p2, region),
            "<<Price 3>>": format_price(p3, region),
            "<<Price 4>>": format_price(gst_amount, region),
            "<<Price 5>>": format_price(total_amount, region),
            "<<Price 6>>": format_price(p1, region),
            "<<Price 7>>": format_price(p2, region),
            "<<Price 8>>": format_price(p3, region),
        })

    elif payment_option == "5 EMI":
        template_name = f"Invoice Template - {region} - 5 EMI Payment Schedule 1.docx"
        p1 = round(total_amount * 0.20)
        p2 = round(total_amount * 0.20)
        p3 = round(total_amount * 0.20)
        p4 = round(total_amount * 0.20)
        p5 = total_amount - (p1 + p2 + p3 + p4)
        placeholders.update({
            "<<Price 1>>": format_price(p1, region),
            "<<Price 2>>": format_price(p2, region),
            "<<Price 3>>": format_price(p3, region),
            "<<Price 4>>": format_price(p4, region),
            "<<Price 5>>": format_price(p5, region),
            "<<Price 6>>": format_price(p1, region),
            "<<Price 7>>": format_price(p2, region),
            "<<Price 8>>": format_price(p3, region),
            "<<Price 9>>": format_price(p4, region),
            "<<Price 10>>": format_price(p5, region),
            "<<Total 1>>": format_price(total_amount, region),
        })

    # Generate Invoice Button
    if st.button("Generate Invoice"):
        try:
            # Clear previous session state data
            if 'invoice_docx' in st.session_state:
                del st.session_state.invoice_docx
                del st.session_state.invoice_pdf
                del st.session_state.invoice_docx_name
                del st.session_state.invoice_pdf_name

            # Generate the next invoice number
            invoice_number = get_next_invoice_number()
            placeholders["<<Invoice>>"] = str(invoice_number)
            placeholders["<<Invoice No>>"] = str(invoice_number)

            # Define the invoice template file path
            template_path = os.path.join(os.getcwd(), template_name)

            # Save the invoice to a temporary directory
            temp_dir = tempfile.gettempdir()
            sanitized_client_name = "".join([c if c.isalnum() or c.isspace() else "_" for c in client_name])
            docx_output_path = os.path.join(temp_dir, f"Invoice_{sanitized_client_name}_{invoice_number}.docx")

            # Edit the template and save the invoice
            edit_invoice_template(template_path, docx_output_path, placeholders)
            
            # Read DOCX file bytes into session state
            with open(docx_output_path, "rb") as file:
                st.session_state.invoice_docx = file.read()
            st.session_state.invoice_docx_name = f"Invoice_{sanitized_client_name}_{invoice_number}.docx"

            # Generate PDF and store in session state
            pdf_output_path = os.path.join(temp_dir, f"Invoice_{sanitized_client_name}_{invoice_number}.pdf")
            pythoncom.CoInitialize()  # Initialize COM library
            try:
                convert(docx_output_path, pdf_output_path)
                with open(pdf_output_path, "rb") as file:
                    st.session_state.invoice_pdf = file.read()
                st.session_state.invoice_pdf_name = f"Invoice_{sanitized_client_name}_{invoice_number}.pdf"
            finally:
                pythoncom.CoUninitialize()

            st.success(f"✅ Invoice #{invoice_number} generated successfully!")

        except Exception as e:
            st.error(f"An error occurred: {e}")

    # Display download buttons if data exists in session state
    if 'invoice_docx' in st.session_state and 'invoice_pdf' in st.session_state:
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label="📥 Download Invoice (Word)",
                data=st.session_state.invoice_docx,
                file_name=st.session_state.invoice_docx_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        with col2:
            st.download_button(
                label="📥 Download Invoice (PDF)",
                data=st.session_state.invoice_pdf,
                file_name=st.session_state.invoice_pdf_name,
                mime="application/pdf"
            )
            
def convert_to_pdf(input_path, output_path):
    """Convert DOCX to PDF safely, handling COM initialization."""
    pythoncom.CoInitialize()  # Initialize COM library
    try:
        convert(input_path, output_path)
    except Exception as e:
        raise Exception(f"PDF conversion error: {e}")
    finally:
        pythoncom.CoUninitialize()  # Uninitialize COM library


# Main App
def main():
    st.sidebar.title("Select Application")
    app_choice = st.sidebar.radio("Choose an application:", ["NDA", "Contract", "Invoice", "Pricing List", "Proposal","Hiring Contract"])
    if app_choice == "NDA" or app_choice=="Contract":
        option = st.selectbox("Select Document Type", ["NDA", "Contract"], key="doc_type")
        generate_document(option)

    elif app_choice == "Invoice":
        generate_invoice()

    elif app_choice == "Hiring Contract":
        generate_hiring_contract()

if __name__ == "__main__":
    main()
