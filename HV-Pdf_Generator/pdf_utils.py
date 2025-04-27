import os
import platform
import subprocess
import shutil
import tempfile
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

logger = logging.getLogger("pdf_utils")

# ========== FORMATTING HELPERS ==========

def apply_formatting(run, font_name="Calibri", font_size=11, bold=False):
    """Apply specific formatting to a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold

def apply_image_placeholder(doc, placeholder_key, image_file):
    """Replace a placeholder with an image."""
    placeholder_found = False

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder_key in para.text:
                        para.clear()
                        run = para.add_run()
                        run.add_picture(image_file, width=Inches(1.5), height=Inches(0.75))
                        placeholder_found = True

    for para in doc.paragraphs:
        if placeholder_key in para.text:
            para.clear()
            run = para.add_run()
            run.add_picture(image_file, width=Inches(1.2), height=Inches(0.75))
            placeholder_found = True

    if not placeholder_found:
        logger.warning(f"Placeholder '{placeholder_key}' not found in the document.")

    return doc

# ========== DOCX -> PDF CONVERTER ==========

def convert_to_pdf(doc_path, pdf_path):
    """Convert Word Document to PDF."""
    doc_path = os.path.abspath(doc_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found at {doc_path}")

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_pdf_path = os.path.join(temp_dir, "temp_output.pdf")

        if platform.system() == "Windows":
            try:
                import comtypes.client
                import pythoncom
                pythoncom.CoInitialize()
                word = comtypes.client.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(doc_path)
                doc.SaveAs(temp_pdf_path, FileFormat=17)  # 17 = PDF
                doc.Close()
                word.Quit()
                pythoncom.CoUninitialize()
            except Exception as e:
                raise Exception(f"Error using COM on Windows: {e}")
        else:
            try:
                subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, doc_path],
                    check=True
                )
                temp_pdf_path = os.path.join(temp_dir, os.path.basename(doc_path).replace('.docx', '.pdf'))
            except subprocess.CalledProcessError as e:
                raise Exception(f"Error using LibreOffice: {e}")

        shutil.copy(temp_pdf_path, pdf_path)

        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"Flattened PDF file was not saved correctly: {pdf_path}")
