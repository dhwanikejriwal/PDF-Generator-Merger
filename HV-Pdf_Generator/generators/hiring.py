import streamlit as st
import os
import tempfile
import uuid
from docx import Document
import fitz  # PyMuPDF
from PIL import Image
import locale
import subprocess
import platform


# Set locale for number formatting
locale.setlocale(locale.LC_ALL, '')

# ---- Function to Replace Placeholders in Word Document ----
def replace_text_in_paragraph(paragraph, placeholders):
    """Replace placeholders in a paragraph."""
    if not paragraph.runs:
        return

    full_text = ''.join(run.text for run in paragraph.runs)

    for key, value in placeholders.items():
        if key in full_text:
            full_text = full_text.replace(key, value)

    # Remove all runs except first
    for i in range(len(paragraph.runs) - 1, 0, -1):
        paragraph._p.remove(paragraph.runs[i]._r)

    if paragraph.runs:
        paragraph.runs[0].text = full_text

def replace_text_in_table(table, placeholders):
    """Replace placeholders inside all cells of a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, placeholders)

def edit_hiring_template(template_path, output_path, placeholders):
    """Edit hiring contract template and save filled version."""
    doc = Document(template_path)

    for para in doc.paragraphs:
        replace_text_in_paragraph(para, placeholders)

    for table in doc.tables:
        replace_text_in_table(table, placeholders)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

# ---- Format price with commas ----
def format_price_with_commas(price_str):
    """Format price string with commas for thousands."""
    try:
        # Remove any existing commas and spaces
        price_str = price_str.replace(',', '').replace(' ', '')
        # Convert to float then back to string with commas
        price_float = float(price_str)
        # Format with comma as thousand separator
        return f"₹{price_float:,.2f}".rstrip('0').rstrip('.') if '.' in f"{price_float:,.2f}" else f"₹{price_float:,.0f}"
    except ValueError:
        # If conversion fails, return original string
        return price_str

# ---- Function to Convert Word to PDF using alternative methods ----
import subprocess
import os
import platform
import streamlit as st

def convert_word_to_pdf(word_path, pdf_path):
    try:
        # First, try LibreOffice (best option)
        if platform.system() == "Windows":
            libreoffice_paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
            
            for lo_path in libreoffice_paths:
                if os.path.exists(lo_path):
                    cmd = [lo_path, '--headless', '--convert-to', 'pdf', '--outdir',
                           os.path.dirname(pdf_path), word_path]
                    subprocess.run(cmd, check=True, timeout=30)
                    
                    # Rename if needed
                    lo_pdf = os.path.splitext(word_path)[0] + '.pdf'
                    if lo_pdf != pdf_path and os.path.exists(lo_pdf):
                        os.rename(lo_pdf, pdf_path)

                    if os.path.exists(pdf_path):
                        return True

        else:
            # Linux / Mac (Streamlit Cloud is Linux)
            cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
                   os.path.dirname(pdf_path), word_path]
            subprocess.run(cmd, check=True, timeout=30)
            
            lo_pdf = os.path.splitext(word_path)[0] + '.pdf'
            if lo_pdf != pdf_path and os.path.exists(lo_pdf):
                os.rename(lo_pdf, pdf_path)

            if os.path.exists(pdf_path):
                return True

        # If LibreOffice not available, try unoconv
        try:
            subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_path, word_path], check=True, timeout=30)
            if os.path.exists(pdf_path):
                return True
        except Exception as e:
            st.warning(f"unoconv method failed: {e}")

        # If everything fails
        st.error("PDF conversion failed. Word document is still available for download.")
        return False

    except Exception as e:
        st.error(f"Unexpected error during Word to PDF conversion: {e}")
        return False


# ---- Function to Render PDF Page ----
def render_pdf_page(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        page = doc[0]  # First page
        pix = page.get_pixmap(dpi=300)  # Higher resolution
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        doc.close()
        return img
    except Exception as e:
        st.error(f"Error rendering PDF: {e}")
        return None

# ---- Navigation Functions ----
def next_page():
    st.session_state.page += 1
    st.rerun()

def prev_page():
    st.session_state.page -= 1
    st.rerun()

# ---- Main App ----
def generate_hiring():
    # Template paths
    template_word = "Hiring Contract.docx"
    
    # Initialize page state
    if "page" not in st.session_state:
        st.session_state.page = 1

    # Page 1: Form Input
    if st.session_state.page == 1:
        st.title("Hiring Document Generator")
        st.subheader("Fill in the Candidate Details")

        with st.form("Hiring Form"):
            today = st.date_input("Today's Date")
            name = st.text_input("Candidate Name")
            role = st.text_input("Internship Role")
            starting_date = st.date_input("Internship Starting Date")
            stipend = st.text_input("Monthly Stipend (in Rs.)")
            working_hours = st.text_input("Working Hours per Week")
            internship_duration = st.text_input("Internship Duration (in months)")
            first_pay = st.date_input("First Pay Date")

            submitted = st.form_submit_button("Generate Hiring Contract")

            if submitted:
                # Format dates nicely
                formatted_today = today.strftime("%d %B, %Y")
                formatted_starting = starting_date.strftime("%d %B, %Y")
                formatted_first_pay = first_pay.strftime("%d %B, %Y")
                
                # Format stipend with commas
                formatted_stipend = format_price_with_commas(stipend)

                # Create replacements dictionary
                replacements = {
                    "<<Date>>": formatted_today,
                    "<<Name>>": name,
                    "<<Role>>": role,
                    "<<Starting Date>>": formatted_starting,
                    "<<Stipend>>": formatted_stipend,
                    "<<Working Hours>>": working_hours,
                    "<<Internship Duration>>": internship_duration,
                    "<<First Pay>>": formatted_first_pay,
                    "<<Contact Email>>": "info@hvtechnologies.app & hvtechnologies19@gmail.com"
                }
                
                # Create temporary filenames with Name-Role format
                temp_dir = tempfile.gettempdir()
                
                # Create a sanitized filename format: "Name-Role Offer Letter"
                sanitized_name = "".join(c for c in name if c.isalnum() or c in [' ', '_']).strip()
                sanitized_role = "".join(c for c in role if c.isalnum() or c in [' ', '_']).strip()
                file_prefix = f"{sanitized_name}-{sanitized_role} Offer Letter"
                sanitized_file_prefix = file_prefix.replace(' ', '_')
                
                filled_word = os.path.join(temp_dir, f"{sanitized_file_prefix}.docx")
                filled_pdf = os.path.join(temp_dir, f"{sanitized_file_prefix}.pdf")
                
                # Generate Word document
                with st.spinner("Generating your document..."):
                    word_success = edit_hiring_template(template_word, filled_word, replacements)
                    
                    # Convert to PDF if Word was successful
                    if word_success:
                        pdf_success = convert_word_to_pdf(filled_word, filled_pdf)
                    else:
                        pdf_success = False
                
                # Store paths and info in session state
                if word_success:
                    st.session_state.filled_word = filled_word
                    st.session_state.candidate_name = name
                    st.session_state.role_name = role
                    st.session_state.file_prefix = file_prefix
                    
                    if pdf_success:
                        st.session_state.filled_pdf = filled_pdf
                    
                    st.success("Document generated successfully!")
                    next_page()
                else:
                    st.error("Failed to generate document. Please try again.")

    # Page 2: Document Preview
    elif st.session_state.page == 2:
        st.title("Hiring Document Preview")
        
        if "filled_pdf" not in st.session_state and "filled_word" not in st.session_state:
            st.warning("No filled document found. Please start again.")
            if st.button("Start Again"):
                st.session_state.page = 1
                st.rerun()
        else:
            # Add container for preview
            preview_container = st.container()
            with preview_container:
                if "filled_pdf" in st.session_state:
                    # Show PDF preview
                    pdf_image = render_pdf_page(st.session_state.filled_pdf)
                    if pdf_image:
                        st.image(pdf_image, use_column_width=True)
                    else:
                        st.warning("Couldn't preview the PDF document.")
                else:
                    st.info("PDF preview not available, but Word document has been generated.")

            # Navigation buttons
            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("Previous"):
                    prev_page()
            with col2:
                if st.button("Next"):
                    next_page()

    # Page 3: Download Options
    elif st.session_state.page == 3:
        st.title("Download Your Hiring Documents 📥")
        
        file_prefix = st.session_state.get("file_prefix", "Offer_Letter")
        
        # Word Download button
        if "filled_word" in st.session_state:
            with open(st.session_state.filled_word, "rb") as f:
                st.download_button(
                    label="Download as Word",
                    data=f,
                    file_name=f"{file_prefix}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # PDF Download button (if available)
        if "filled_pdf" in st.session_state:
            with open(st.session_state.filled_pdf, "rb") as f:
                st.download_button(
                    label="Download as PDF",
                    data=f,
                    file_name=f"{file_prefix}.pdf",
                    mime="application/pdf"
                )

        # New document button
        if st.button("Create Another Hiring Contract"):
            st.session_state.page = 1
            st.rerun()