import streamlit as st
from datetime import datetime
import os
from docx import Document
from num2words import num2words

from pdf_utils import convert_to_pdf
from session_manager import clear_session_keys

# ========== Helper Functions ==========

def format_price(amount, currency):
    """Format price based on currency."""
    formatted_price = f"{amount:,.2f}"
    return f"{currency} {formatted_price}" if currency == "USD" else f"Rs. {formatted_price}"

def amount_to_words(amount):
    """Convert amount to words (English)."""
    try:
        words = num2words(amount, lang='en').replace(',', '').title()
        return words
    except Exception:
        return f"[Error converting {amount}]"

def replace_placeholders(doc, placeholders):
    """Replace placeholders in paragraphs and tables."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        if key.startswith("<<Price") or key.startswith("<<Total") or key == "<<Amt to word>>":
                            run.bold = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, value in placeholders.items():
                            if key in run.text:
                                run.text = run.text.replace(key, value)
                                if key.startswith("<<Price") or key.startswith("<<Total") or key == "<<Amt to word>>":
                                    run.bold = True
    return doc

def edit_invoice_template(template_path, output_path, placeholders):
    """Edit invoice template and save filled version."""
    doc = Document(template_path)
    replace_placeholders(doc, placeholders)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

def get_next_invoice_number():
    """Simple invoice number counter stored in file."""
    invoice_file = "invoice_counter.txt"
    if not os.path.exists(invoice_file):
        with open(invoice_file, "w") as f:
            f.write("1000")  # Start with 1000
    try:
        with open(invoice_file, "r") as f:
            current_invoice = int(f.read().strip())
    except ValueError:
        current_invoice = 1000
    next_invoice = current_invoice + 1
    with open(invoice_file, "w") as f:
        f.write(str(next_invoice))
    return current_invoice

# ========== Main Generator Function ==========

def generate_invoice():
    st.title("Invoice Generator")

    region = st.selectbox("Select Region", ["INR", "USD"])
    payment_options = ["1 Payment", "3 EMI", "5 EMI"] if region == "INR" else ["3 EMI", "5 EMI"]

    client_name = st.text_input("Client Name")
    client_address = st.text_input("Client Address")
    project_name = st.text_input("Project Name")
    phone_number = st.text_input("Phone Number")
    gst_number = st.text_input("GST Number")
    base_amount = st.number_input("Base Amount (excluding GST)", min_value=0.0, format="%.2f")
    payment_option = st.selectbox("Payment Option", payment_options)
    invoice_date = st.date_input("Invoice Date", value=datetime.today())

    gst_amount = round(base_amount * 0.18)
    total_amount = base_amount + gst_amount

    placeholders = {
        "<<Client Name>>": client_name,
        "<<Client Address>>": client_address,
        "<<GST Number>>": gst_number,
        "<<Client Email>>": client_address,
        "<<Project Name>>": project_name,
        "<<Mobile Number>>": phone_number,
        "<<Date>>": invoice_date.strftime("%d-%m-%Y"),
        "<<Amt to word>>": amount_to_words(int(total_amount)),
    }

    # Select template based on payment option
    if payment_option == "1 Payment":
        template_name = f"Invoice Template - {region} - 1 Payment 1.docx"
        placeholders.update({
            "<<Price 1>>": format_price(base_amount, region),
            "<<Price 2>>": format_price(gst_amount, region),
            "<<Price 3>>": format_price(total_amount, region),
            "<<Total 1>>": format_price(total_amount, region),
        })
    elif payment_option == "3 EMI":
        template_name = f"Invoice Template - {region} - 3 EMI Payment Schedule 1.docx"
        p1 = round(total_amount * 0.30)
        p2 = round(total_amount * 0.40)
        p3 = total_amount - (p1 + p2)
        placeholders.update({
            "<<Price 1>>": format_price(p1, region),
            "<<Price 2>>": format_price(p2, region),
            "<<Price 3>>": format_price(p3, region),
            "<<Price 4>>": format_price(gst_amount, region),
            "<<Price 5>>": format_price(total_amount, region),
            "<<Price 6>>": format_price(p1, region),
            "<<Price 7>>": format_price(p2, region),
            "<<Price 8>>": format_price(p3, region),
        })
    elif payment_option == "5 EMI":
        template_name = f"Invoice Template - {region} - 5 EMI Payment Schedule 1.docx"
        p1 = round(total_amount * 0.20)
        p2 = round(total_amount * 0.20)
        p3 = round(total_amount * 0.20)
        p4 = round(total_amount * 0.20)
        p5 = total_amount - (p1 + p2 + p3 + p4)
        placeholders.update({
            "<<Price 1>>": format_price(p1, region),
            "<<Price 2>>": format_price(p2, region),
            "<<Price 3>>": format_price(p3, region),
            "<<Price 4>>": format_price(p4, region),
            "<<Price 5>>": format_price(p5, region),
            "<<Price 6>>": format_price(p1, region),
            "<<Price 7>>": format_price(p2, region),
            "<<Price 8>>": format_price(p3, region),
            "<<Price 9>>": format_price(p4, region),
            "<<Price 10>>": format_price(p5, region),
            "<<Total 1>>": format_price(total_amount, region),
        })

    output_dir = os.path.join("app", "generated_files", "invoices")
    os.makedirs(output_dir, exist_ok=True)

    if st.button("Generate Invoice"):
        try:
            clear_session_keys(["invoice_docx", "invoice_pdf", "invoice_docx_name", "invoice_pdf_name"])

            invoice_number = get_next_invoice_number()
            placeholders["<<Invoice>>"] = str(invoice_number)
            placeholders["<<Invoice No>>"] = str(invoice_number)

            template_path = os.path.join(os.getcwd(), template_name)

            if not os.path.exists(template_path):
                st.error(f"Template file not found: {template_path}")
                return

            safe_client_name = ''.join(c if c.isalnum() else '_' for c in client_name)

            docx_output_path = os.path.join(output_dir, f"Invoice_{safe_client_name}_{invoice_number}.docx")
            pdf_output_path = os.path.join(output_dir, f"Invoice_{safe_client_name}_{invoice_number}.pdf")

            # Generate DOCX
            edit_invoice_template(template_path, docx_output_path, placeholders)

            # Save DOCX to session
            with open(docx_output_path, "rb") as docx_file:
                st.session_state.invoice_docx = docx_file.read()
                st.session_state.invoice_docx_name = f"Invoice_{safe_client_name}_{invoice_number}.docx"

            # Convert to PDF
            try:
                convert_to_pdf(docx_output_path, pdf_output_path)

                if os.path.exists(pdf_output_path):
                    with open(pdf_output_path, "rb") as pdf_file:
                        st.session_state.invoice_pdf = pdf_file.read()
                        st.session_state.invoice_pdf_name = f"Invoice_{safe_client_name}_{invoice_number}.pdf"
                else:
                    st.warning("PDF not found after conversion.")
            except Exception as pdf_err:
                st.error(f"PDF Conversion Error: {pdf_err}")
                st.warning("PDF conversion failed, but DOCX is available.")

            # Download buttons
            col1, col2 = st.columns(2)

            with col1:
                if st.session_state.invoice_docx:
                    st.download_button(
                        label="📥 Download Invoice (Word)",
                        data=st.session_state.invoice_docx,
                        file_name=st.session_state.invoice_docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            with col2:
                if st.session_state.invoice_pdf:
                    st.download_button(
                        label="📥 Download Invoice (PDF)",
                        data=st.session_state.invoice_pdf,
                        file_name=st.session_state.invoice_pdf_name,
                        mime="application/pdf"
                    )

        except Exception as e:
            st.error(f"An error occurred: {e}")
            import traceback
            st.code(traceback.format_exc())
