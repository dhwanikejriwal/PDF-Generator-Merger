import streamlit as st
import os
import tempfile
import uuid
from docx import Document
import fitz  # PyMuPDF
from PIL import Image
import locale
import subprocess
import platform
import datetime
def replace_text_in_paragraph(paragraph, placeholders):
    """Replace placeholders in a paragraph, preserving formatting and optionally bolding specific runs."""
    # Combine all run texts
    full_text = ''.join(run.text for run in paragraph.runs)

    for key, value in placeholders.items():
        if key in full_text:
            before, sep, after = full_text.partition(key)
            # Remove existing runs
            for run in list(paragraph.runs)[::-1]:
                paragraph._p.remove(run._r)
            # Add text before placeholder
            if before:
                paragraph.add_run(before)
            # Add replaced text
            new_run = paragraph.add_run(value)
            # Uncomment the next lines if you want to bold any specific placeholders:
            # if key in ('<<ClientName>>', '<<Date>>'):
            #     new_run.font.bold = True
            # Add text after placeholder
            if after:
                paragraph.add_run(after)
            # Stop after first replacement in this paragraph
            return


def replace_text_in_table(table, placeholders):
    """Iterate through table cells and replace placeholders."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, placeholders)


def edit_nda_template(template_path, output_path, placeholders):
    """Load template, replace placeholders in body and tables, then save."""
    doc = Document(template_path)

    # Replace in document body
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, placeholders)

    # Replace in all tables
    for table in doc.tables:
        replace_text_in_table(table, placeholders)

    # Ensure target directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_nda():
    """Streamlit UI to collect inputs and generate the NDA document."""
    st.title("NDA Generator")

    client_name = st.text_input("Enter Client Name:")
    company_name = st.text_input("Enter Company Name:")
    date_input = st.date_input("Enter Date:", datetime.today())
    address = st.text_area("Enter Address:")

    placeholders = {
        "<<ClientName>>": client_name,
        "<<CompanyName>>": company_name,
        "<<Date>>": date_input.strftime("%d-%m-%Y"),
        "<<Address>>": address,
    }

    template_name = "NDA Template.docx"
    temp_dir = tempfile.gettempdir()


    if st.button("Generate NDA"):
        try:
            # Clear previous session state data
            for key in ["nda_docx", "nda_pdf", "nda_docx_name", "nda_pdf_name"]:
                if key in st.session_state:
                    st.session_state[key] = None if "name" not in key else ""

            # Define the hiring template file path
            template_path = os.path.join(os.getcwd(), template_name)
            
            # Verify the template exists
            if not os.path.exists(template_path):
                st.error(f"Template file not found: {template_path}")
                return

            
            safe_name = ''.join(c if c.isalnum() else '_' for c in client_name)
            
            # Save the hiring contract to a temporary directory
            temp_dir = tempfile.gettempdir()
            docx_output_path = os.path.join(temp_dir, f"NDA_{safe_name}.docx")
            pdf_output_path = os.path.join(temp_dir, f"NDA_{safe_name}.pdf")

            # Edit the hiring template and save the contract
            edit_nda_template(template_path, docx_output_path, placeholders)
            # st.info("DOCX file created successfully. Converting to PDF...")

            # Load the generated DOCX file into session state for download
            with open(docx_output_path, "rb") as docx_file:
                st.session_state.nda_docx = docx_file.read()
                st.session_state.nda_docx_name = f"NDA_{safe_name}.docx"

            # Convert DOCX to PDF with better error handling
            try:
                convert_to_pdf(docx_output_path, pdf_output_path)
                # st.info(f"PDF conversion completed. Checking result...")
                
                if os.path.exists(pdf_output_path):
                    with open(pdf_output_path, "rb") as pdf_file:
                        st.session_state.nda_pdf = pdf_file.read()
                        st.session_state.nda_pdf_name =f"NDA_{safe_name}.pdf"
                    # st.success("PDF created successfully!")
                else:
                    st.warning("PDF file not found after conversion attempt.")
            except Exception as pdf_err:
                st.error(f"PDF Conversion Error: {pdf_err}")
                # Still allow DOCX download even if PDF fails
                st.warning("PDF conversion failed, but DOCX is available for download.")

            # Display download buttons based on what's available
            col1, col2 = st.columns(2)
            
            with col1:
                if st.session_state.nda_docx:
                    st.download_button(
                        label="📥 Download NDA(Word)",
                        data=st.session_state.nda_docx,
                        file_name=st.session_state.nda_docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.warning("DOCX file not available for download.")
                    
            with col2:
                if st.session_state.nda_pdf:
                    st.download_button(
                        label="📥 Download NDA(PDF)",
                        data=st.session_state.nda_pdf,
                        file_name=st.session_state.nda_pdf_name,
                        mime="application/pdf"
                    )
                else:
                    st.warning("PDF file not available for download.")
                    
        except Exception as e:
            st.error(f"An error occurred: {e}")
            import traceback
            st.code(traceback.format_exc())
