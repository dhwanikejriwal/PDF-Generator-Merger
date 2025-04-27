import streamlit as st
from datetime import datetime
import os
from docx import Document

from pdf_utils import convert_to_pdf
from session_manager import clear_session_keys

# ========== Helper Functions ==========

def replace_placeholders(doc, placeholders):
    """Replace placeholders in paragraphs and tables."""
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        if key == "<<EndDate>>":
                            run.bold = True  # Bold specific keys

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for key, value in placeholders.items():
                            if key in run.text:
                                run.text = run.text.replace(key, value)
                                if key == "<<EndDate>>":
                                    run.bold = True
    return doc

def edit_contract_template(template_path, output_path, placeholders):
    """Edit contract template and save filled version."""
    doc = Document(template_path)
    replace_placeholders(doc, placeholders)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

# ========== Main Generator Function ==========

def generate_contract():
    st.title("Contract Generator")

    client_name = st.text_input("Enter Client Name:")
    company_name = st.text_input("Enter Company Name:")
    date_input = st.date_input("Enter Effective Date:", datetime.today())
    end_date_input = st.date_input("Enter Contract End Date:", datetime.today())
    address = st.text_area("Enter Address:")

    placeholders = {
        "<<ClientName>>": client_name,
        "<<CompanyName>>": company_name,
        "<<Date>>": date_input.strftime("%d-%m-%Y"),
        "<<StartDate>>": date_input.strftime("%d %B %Y"),
        "<<EndDate>>": end_date_input.strftime("%d %B %Y"),
        "<<Address>>": address,
    }

    template_name = "Contract Template.docx"
    output_dir = os.path.join("app", "generated_files", "contracts")
    os.makedirs(output_dir, exist_ok=True)

    if st.button("Generate Contract"):
        try:
            clear_session_keys(["contract_docx", "contract_pdf", "contract_docx_name", "contract_pdf_name"])

            template_path = os.path.join(os.getcwd(), template_name)

            if not os.path.exists(template_path):
                st.error(f"Template file not found: {template_path}")
                return

            safe_name = ''.join(c if c.isalnum() else '_' for c in client_name)

            docx_output_path = os.path.join(output_dir, f"Contract_{safe_name}.docx")
            pdf_output_path = os.path.join(output_dir, f"Contract_{safe_name}.pdf")

            # Generate DOCX
            edit_contract_template(template_path, docx_output_path, placeholders)

            # Save DOCX to session
            with open(docx_output_path, "rb") as docx_file:
                st.session_state.contract_docx = docx_file.read()
                st.session_state.contract_docx_name = f"Contract_{safe_name}.docx"

            # Convert to PDF
            try:
                convert_to_pdf(docx_output_path, pdf_output_path)

                if os.path.exists(pdf_output_path):
                    with open(pdf_output_path, "rb") as pdf_file:
                        st.session_state.contract_pdf = pdf_file.read()
                        st.session_state.contract_pdf_name = f"Contract_{safe_name}.pdf"
                else:
                    st.warning("PDF not found after conversion.")
            except Exception as pdf_err:
                st.error(f"PDF Conversion Error: {pdf_err}")
                st.warning("PDF conversion failed, but DOCX is available.")

            # Download buttons
            col1, col2 = st.columns(2)

            with col1:
                if st.session_state.contract_docx:
                    st.download_button(
                        label="📥 Download Contract (Word)",
                        data=st.session_state.contract_docx,
                        file_name=st.session_state.contract_docx_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            with col2:
                if st.session_state.contract_pdf:
                    st.download_button(
                        label="📥 Download Contract (PDF)",
                        data=st.session_state.contract_pdf,
                        file_name=st.session_state.contract_pdf_name,
                        mime="application/pdf"
                    )

        except Exception as e:
            st.error(f"An error occurred: {e}")
            import traceback
            st.code(traceback.format_exc())
